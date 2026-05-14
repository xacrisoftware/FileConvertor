import os


DOCUMENT_FORMATS = {"TXT", "PDF", "DOCX"}

EXT_MAP = {
    ".txt": "TXT",
    ".pdf": "PDF",
    ".docx": "DOCX",
}


class DocumentConvertError(Exception):
    pass


def get_format_from_ext(path):
    ext = os.path.splitext(path)[1].lower()
    return EXT_MAP.get(ext)


def convert_document(input_path, output_path, output_format):
    in_fmt = get_format_from_ext(input_path)
    out_fmt = output_format.upper()

    if not in_fmt:
        raise DocumentConvertError("Неизвестный формат входного файла")

    if in_fmt == out_fmt:
        raise DocumentConvertError("Входной и выходной формат совпадают")

    if out_fmt not in DOCUMENT_FORMATS:
        raise DocumentConvertError(f"Неподдерживаемый формат: {output_format}")

    conv_key = f"{in_fmt}_TO_{out_fmt}"

    handlers = {
        "TXT_TO_PDF": _txt_to_pdf,
        "TXT_TO_DOCX": _txt_to_docx,
        "DOCX_TO_TXT": _docx_to_txt,
        "PDF_TO_TXT": _pdf_to_txt,
    }

    handler = handlers.get(conv_key)
    if not handler:
        raise DocumentConvertError(f"Конвертация {in_fmt} → {out_fmt} не поддерживается")

    return handler(input_path, output_path)


def _find_unicode_font(pdf):
    win_fonts = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts")
    candidates = [
        ("Arial", os.path.join(win_fonts, "arial.ttf")),
        ("SegoeUI", os.path.join(win_fonts, "segoeui.ttf")),
        ("Calibri", os.path.join(win_fonts, "calibri.ttf")),
        ("Times", os.path.join(win_fonts, "times.ttf")),
    ]
    for name, path in candidates:
        if os.path.exists(path):
            try:
                pdf.add_font(name, "", path, uni=True)
                pdf.set_font(name, "", 12)
                return
            except Exception:
                continue
    pdf.set_font("Courier", "", 12)


def _txt_to_pdf(input_path, output_path):
    try:
        from fpdf import FPDF
    except ImportError:
        raise DocumentConvertError("Установите fpdf2: pip install fpdf2")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    _find_unicode_font(pdf)

    try:
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(input_path, "r", encoding="cp1251") as f:
            text = f.read()

    for line in text.split("\n"):
        try:
            pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")
        except Exception:
            pdf.cell(0, 6, "[encoding error]", new_x="LMARGIN", new_y="NEXT")

    pdf.output(output_path)
    return output_path


def _txt_to_docx(input_path, output_path):
    try:
        from docx import Document
    except ImportError:
        raise DocumentConvertError("Установите python-docx: pip install python-docx")

    doc = Document()
    try:
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(input_path, "r", encoding="cp1251") as f:
            text = f.read()

    for para in text.split("\n"):
        doc.add_paragraph(para)

    doc.save(output_path)
    return output_path


def _docx_to_txt(input_path, output_path):
    try:
        from docx import Document
    except ImportError:
        raise DocumentConvertError("Установите python-docx: pip install python-docx")

    doc = Document(input_path)
    lines = [p.text for p in doc.paragraphs]

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return output_path


def _pdf_to_txt(input_path, output_path):
    try:
        import fitz
    except ImportError:
        try:
            import pdfplumber
        except ImportError:
            raise DocumentConvertError(
                "Установите PyMuPDF (fitz) или pdfplumber: pip install pymupdf"
            )
        return _pdf_to_txt_plumber(input_path, output_path)

    doc = fitz.open(input_path)
    lines = []
    for page in doc:
        lines.append(page.get_text())
    doc.close()

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return output_path


def _pdf_to_txt_plumber(input_path, output_path):
    import pdfplumber

    lines = []
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return output_path
